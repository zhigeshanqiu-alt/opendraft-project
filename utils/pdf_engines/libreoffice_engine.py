#!/usr/bin/env python3
"""
ABOUTME: LibreOffice-based PDF generation engine using headless conversion
ABOUTME: Generates DOCX first, then converts to PDF via LibreOffice CLI
"""

import re
import subprocess
import shutil
from pathlib import Path
from typing import Optional

from .base import PDFEngine, PDFGenerationOptions, EngineResult

# Import DOCX generation from python-docx
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class LibreOfficeEngine(PDFEngine):
    """
    LibreOffice-based PDF generation engine.

    Two-stage process:
    1. Generate DOCX from markdown using python-docx
    2. Convert DOCX to PDF using LibreOffice headless mode

    Advantages:
    - Better font rendering than WeasyPrint
    - Handles typography properly (AI vs Al visual distinction)
    - Academic-quality output
    - Moderate installation size (~200MB)

    Requirements:
    - python-docx library
    - libreoffice-writer (headless)
    """

    def get_name(self) -> str:
        """Get engine name."""
        return "LibreOffice"

    def get_priority(self) -> int:
        """
        Priority: 70/100

        Higher than WeasyPrint (better rendering), lower than Pandoc/LaTeX
        (less academic standard).
        """
        return 70

    def is_available(self) -> bool:
        """Check if python-docx and libreoffice are available."""
        if not DOCX_AVAILABLE:
            return False

        # Check if libreoffice command exists (libreoffice on Linux, soffice on macOS)
        return shutil.which('libreoffice') is not None or shutil.which('soffice') is not None

    def _get_libreoffice_cmd(self) -> str:
        """Get the correct LibreOffice command for this platform."""
        if shutil.which('libreoffice'):
            return 'libreoffice'
        return 'soffice'

    def generate(
        self,
        md_file: Path,
        output_pdf: Path,
        options: PDFGenerationOptions
    ) -> EngineResult:
        """
        Generate PDF via DOCX → LibreOffice conversion.

        Args:
            md_file: Input markdown file
            output_pdf: Output PDF path
            options: Generation options

        Returns:
            EngineResult with success/failure details
        """
        # Validate inputs
        error = self.validate_inputs(md_file, output_pdf)
        if error:
            return EngineResult(
                success=False,
                engine_name=self.get_name(),
                error_message=error
            )

        try:
            # Step 1: Generate DOCX
            temp_docx = output_pdf.parent / f"{output_pdf.stem}_temp.docx"
            docx_result = self._generate_docx(md_file, temp_docx, options)

            if not docx_result.success:
                return docx_result

            # Step 2: Convert DOCX → PDF using LibreOffice
            pdf_result = self._convert_docx_to_pdf(temp_docx, output_pdf)

            # Cleanup temp file
            if temp_docx.exists():
                temp_docx.unlink()

            return pdf_result

        except Exception as e:
            return EngineResult(
                success=False,
                engine_name=self.get_name(),
                error_message=f"Unexpected error: {str(e)}"
            )

    def _generate_docx(
        self,
        md_file: Path,
        output_docx: Path,
        options: PDFGenerationOptions
    ) -> EngineResult:
        """
        Generate DOCX from markdown using python-docx.

        Args:
            md_file: Input markdown file
            output_docx: Output DOCX path
            options: Generation options

        Returns:
            EngineResult with success/failure
        """
        try:
            # Read markdown
            with open(md_file, 'r', encoding='utf-8') as f:
                lines = f.readlines()

            # Create document
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                # Convert margin string to inches
                margin_value = self._parse_margin(options.margins)
                section.top_margin = margin_value
                section.bottom_margin = margin_value
                section.left_margin = margin_value
                section.right_margin = margin_value

            # Add page numbers to footer
            if options.page_numbers:
                self._add_page_numbers(doc, options.page_number_position)

            # Process markdown content line by line
            for line in lines:
                line = line.rstrip()

                if not line:
                    continue

                # Title (# heading)
                if line.startswith('# '):
                    para = doc.add_heading(line[2:], level=1)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    continue

                # Section heading (## heading)
                if line.startswith('## '):
                    para = doc.add_heading(line[3:], level=2)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    continue

                # Subsection heading (### heading)
                if line.startswith('### '):
                    # Level 3 headings should be italic, not bold (APA 7th edition)
                    para = doc.add_paragraph()
                    run = para.add_run(line[4:])
                    run.italic = True
                    run.bold = False
                    run.font.name = options.font_family
                    run.font.size = Pt(int(options.font_size.replace('pt', '')))

                    # Set paragraph formatting
                    para_format = para.paragraph_format
                    para_format.space_before = Pt(18)
                    para_format.space_after = Pt(6)

                    continue

                # Horizontal rule
                if line.startswith('---'):
                    para = doc.add_paragraph()
                    run = para.add_run('_' * 60)
                    run.font.size = Pt(12)
                    continue

                # Regular paragraph - process inline markdown
                para = doc.add_paragraph()
                self._process_inline_markdown(para, line)

                # Set paragraph formatting
                para_format = para.paragraph_format
                para_format.line_spacing = options.line_spacing
                para_format.space_after = Pt(0)

                # Set font for all runs
                for run in para.runs:
                    run.font.name = options.font_family
                    run.font.size = Pt(int(options.font_size.replace('pt', '')))

            # Save DOCX
            doc.save(output_docx)

            return EngineResult(
                success=True,
                engine_name=f"{self.get_name()} (DOCX generation)",
                output_path=output_docx
            )

        except Exception as e:
            return EngineResult(
                success=False,
                engine_name=f"{self.get_name()} (DOCX generation)",
                error_message=f"Failed to generate DOCX: {str(e)}"
            )

    def _process_inline_markdown(self, paragraph, text: str) -> None:
        """
        Process inline markdown formatting (*italic*, **bold**, `code`, etc.).

        Handles:
        - **bold** text
        - *italic* text
        - ***bold italic*** text
        - `code` spans (monospace)
        - Escaped characters (\\*, \\`, etc.)

        Args:
            paragraph: python-docx paragraph object
            text: Text with inline markdown
        """
        import re

        # Pattern to match markdown inline elements
        # Matches: ***text***, **text**, *text*, `code`
        # In order of priority (longest match first)
        pattern = re.compile(
            r'(\*\*\*(?P<bolditalic>[^\*]+)\*\*\*)|'  # ***bold italic***
            r'(\*\*(?P<bold>[^\*]+)\*\*)|'            # **bold**
            r'(\*(?P<italic>[^\*]+)\*)|'              # *italic*
            r'(`(?P<code>[^`]+)`)'                     # `code`
        )

        current_pos = 0

        for match in pattern.finditer(text):
            # Add text before match (plain text)
            if match.start() > current_pos:
                plain_text = text[current_pos:match.start()]
                paragraph.add_run(plain_text)

            # Add formatted text based on match group
            if match.group('bolditalic'):
                run = paragraph.add_run(match.group('bolditalic'))
                run.bold = True
                run.italic = True
            elif match.group('bold'):
                run = paragraph.add_run(match.group('bold'))
                run.bold = True
            elif match.group('italic'):
                run = paragraph.add_run(match.group('italic'))
                run.italic = True
            elif match.group('code'):
                run = paragraph.add_run(match.group('code'))
                run.font.name = 'Courier New'
                run.font.size = Pt(11)

            current_pos = match.end()

        # Add remaining text after last match
        if current_pos < len(text):
            paragraph.add_run(text[current_pos:])

    def _convert_docx_to_pdf(
        self,
        docx_file: Path,
        output_pdf: Path
    ) -> EngineResult:
        """
        Convert DOCX to PDF using LibreOffice headless mode.

        Args:
            docx_file: Input DOCX file
            output_pdf: Output PDF path

        Returns:
            EngineResult with success/failure
        """
        try:
            # LibreOffice converts to PDF in the same directory as input
            # So we need to handle the output location carefully
            output_dir = output_pdf.parent

            # Run LibreOffice in headless mode
            cmd = [
                self._get_libreoffice_cmd(),
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_dir),
                str(docx_file)
            ]

            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60  # 60 second timeout
            )

            if result.returncode != 0:
                return EngineResult(
                    success=False,
                    engine_name=self.get_name(),
                    error_message=f"LibreOffice conversion failed: {result.stderr}"
                )

            # LibreOffice creates PDF with same basename as DOCX
            generated_pdf = output_dir / f"{docx_file.stem}.pdf"

            # Rename to desired output path if needed
            if generated_pdf != output_pdf and generated_pdf.exists():
                if output_pdf.exists():
                    output_pdf.unlink()
                generated_pdf.rename(output_pdf)

            if not output_pdf.exists():
                return EngineResult(
                    success=False,
                    engine_name=self.get_name(),
                    error_message="LibreOffice did not generate PDF file"
                )

            return EngineResult(
                success=True,
                engine_name=self.get_name(),
                output_path=output_pdf
            )

        except subprocess.TimeoutExpired:
            return EngineResult(
                success=False,
                engine_name=self.get_name(),
                error_message="LibreOffice conversion timed out (>60s)"
            )
        except Exception as e:
            return EngineResult(
                success=False,
                engine_name=self.get_name(),
                error_message=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    def _parse_margin(margin_str: str) -> "Inches":
        """
        Parse margin string to Inches object.

        Args:
            margin_str: Margin string like "1in", "2.5cm", etc.

        Returns:
            Inches object
        """
        # Simple parser for common margin formats
        if margin_str.endswith('in'):
            value = float(margin_str.replace('in', ''))
            return Inches(value)
        elif margin_str.endswith('cm'):
            value = float(margin_str.replace('cm', ''))
            return Inches(value / 2.54)  # Convert cm to inches
        else:
            # Default to 1 inch
            return Inches(1.0)

    @staticmethod
    def _add_page_numbers(doc: "Document", position: str = "bottom-center") -> None:
        """
        Add page numbers to document footer.

        Args:
            doc: Document object
            position: Position string (bottom-center, bottom-right, etc.)
        """
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        # Set alignment based on position
        if 'center' in position:
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'right' in position:
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # Style the page number
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
