#!/usr/bin/env python3
"""
ABOUTME: Post-processor for DOCX files to add academic structure (title page, TOC)
ABOUTME: Fixes Pandoc's inline title block by inserting professional page breaks

Production-grade DOCX post-processing following SOLID principles.
"""

from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH


def insert_academic_structure(
    docx_path: Path,
    verbose: bool = False,
    options: Optional[Dict[str, Any]] = None
) -> bool:
    """
    Insert academic paper structure into Pandoc-generated DOCX.
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    try:
        if verbose:
            print(f"📄 Post-processing DOCX: {docx_path.name}")

        doc = Document(docx_path)

        # Step 1: Find title block elements
        title_idx, date_idx = _find_title_block(doc)

        if title_idx is None:
            if verbose:
                print("   ⚠️  No title block found - skipping post-processing")
            return True

        if verbose:
            print(f"   ✓ Found title block (Title at {title_idx}, Date at {date_idx})")

        # Step 2: Insert institution info BEFORE title
        if options and options.get('institution'):
            _insert_institution_block(doc, title_idx, options, verbose)
            # Recalculate positions after insertion
            title_idx, date_idx = _find_title_block(doc)

        # Step 2b: Center the title block (Title, Subtitle, Author, Date)
        _center_title_block(doc, title_idx, date_idx)

        # Step 3: Insert additional metadata AFTER date (supervisor, etc.)
        if options:
            _insert_metadata_after_date(doc, date_idx, options, verbose)
            # Recalculate date position
            _, date_idx = _find_title_block(doc)

        # Step 4: Find end of cover page and insert page break
        cover_end_idx = _find_cover_end(doc)
        if cover_end_idx is not None:
            _insert_page_break_after(doc, cover_end_idx)
            if verbose:
                print(f"   ✓ Inserted page break after cover page")

        # Step 5: Pandoc generates TOC with --toc flag, so we don't insert manual TOC
        # Just need to insert page break after Abstract (before first chapter)
        abstract_end_idx = _find_abstract_end(doc)
        if abstract_end_idx is not None:
            _insert_page_break_after(doc, abstract_end_idx)
            if verbose:
                print(f"   ✓ Inserted page break after Abstract")

        # Step 6: Fix table widths to fit page
        _fix_table_widths(doc, verbose)

        doc.save(docx_path)

        if verbose:
            print(f"   ✅ Post-processing complete!")

        return True

    except Exception as e:
        if verbose:
            print(f"   ❌ Post-processing failed: {e}")
            import traceback
            traceback.print_exc()
        return False


def _find_title_block(doc: Document):
    """Find Title and Date paragraph indices."""
    title_idx = None
    date_idx = None

    for i, para in enumerate(doc.paragraphs[:20]):
        style = para.style.name if para.style else ""
        if style == 'Title' and title_idx is None:
            title_idx = i
        elif style == 'Date':
            date_idx = i
            break

    return title_idx, date_idx


def _center_title_block(doc: Document, title_idx, date_idx):
    """Center all paragraphs in the title block (Title through Date)."""
    if title_idx is None or date_idx is None:
        return

    for i in range(title_idx, date_idx + 1):
        if i < len(doc.paragraphs):
            doc.paragraphs[i].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _find_cover_end(doc: Document):
    """Find the last paragraph of cover page content (before Abstract or first heading)."""
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text.strip().lower()

        # Cover ends when we hit Abstract or a Heading
        if 'abstract' in text and 'Heading' in style:
            return i - 1
        if style.startswith('Heading') and i > 5:
            return i - 1

    return None


def _find_abstract_heading(doc: Document):
    """Find the Abstract heading paragraph index."""
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text.strip().lower()

        if 'abstract' in text and 'Heading' in style:
            return i

    return None


def _insert_toc(doc: Document, insert_before_idx: int, verbose: bool = False):
    """Insert Table of Contents before the specified index."""
    # Collect all headings from the document
    headings = []
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text.strip()

        if not text:
            continue

        if style == 'Heading 1':
            headings.append((1, text))
        elif style == 'Heading 2':
            headings.append((2, text))
        elif style == 'Heading 3':
            headings.append((3, text))

    if not headings:
        return None

    # Get the paragraph before which to insert
    target_para = doc.paragraphs[insert_before_idx]

    # Insert TOC heading
    toc_heading = target_para.insert_paragraph_before("Table of Contents")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in toc_heading.runs:
        run.font.size = Pt(16)
        run.font.bold = True

    # Insert empty line after heading
    target_para.insert_paragraph_before('')

    # Insert TOC entries
    entry_count = 0
    for level, heading_text in headings:
        indent = "    " * (level - 1)
        entry_para = target_para.insert_paragraph_before(f"{indent}{heading_text}")
        entry_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in entry_para.runs:
            run.font.size = Pt(11)
        entry_count += 1

    # Insert empty line after TOC
    empty_para = target_para.insert_paragraph_before('')

    if verbose:
        print(f"   ✓ Inserted Table of Contents ({entry_count} entries)")

    # Return index of the last TOC paragraph (empty line after entries)
    # We need to find it by searching
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == 'abstract' and 'Heading' in (para.style.name if para.style else ''):
            return i - 1  # Return the paragraph before Abstract

    return insert_before_idx + entry_count + 2  # Approximate


def _find_abstract_end(doc: Document):
    """Find end of Abstract section (before first chapter heading)."""
    in_abstract = False
    last_abstract_para = None

    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text.strip()

        # Start of Abstract
        if 'abstract' in text.lower() and 'Heading' in style:
            in_abstract = True
            continue

        if in_abstract:
            # Check for chapter heading (Heading 1 starting with number)
            if style == 'Heading 1' and text and (text[0].isdigit() or text.startswith('1')):
                return last_abstract_para
            last_abstract_para = i

    return None


def _insert_institution_block(doc: Document, title_idx: int, options: Dict, verbose: bool):
    """Insert institution, faculty, department before title."""
    insert_count = 0

    # Insert in reverse order so indices stay correct
    items = []

    if options.get('department'):
        items.append(('department', options['department'], True, 11))  # italic
    if options.get('faculty'):
        items.append(('faculty', options['faculty'], False, 11))
    if options.get('institution'):
        items.append(('institution', options['institution'].upper(), False, 14))  # uppercase

    # Insert each item at title_idx (they'll stack up correctly)
    for name, text, italic, size in items:
        para = doc.paragraphs[title_idx]
        new_para = para.insert_paragraph_before(text)
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in new_para.runs:
            run.font.size = Pt(size)
            if italic:
                run.font.italic = True
            if name == 'institution':
                run.font.small_caps = True
        insert_count += 1

    # Add empty line after institution block
    if insert_count > 0:
        para = doc.paragraphs[title_idx + insert_count]
        para.insert_paragraph_before('')

    if verbose and insert_count > 0:
        print(f"   ✓ Added institution block ({insert_count} lines)")


def _insert_metadata_after_date(doc: Document, date_idx: int, options: Dict, verbose: bool):
    """Insert supervisor info, student ID, etc. after the date paragraph."""
    if date_idx is None:
        return

    # Find the Date paragraph
    date_para = doc.paragraphs[date_idx]
    insert_after = date_para
    additions = 0

    # Project type
    if options.get('project_type'):
        new_para = _insert_para_after(insert_after, '')
        insert_after = new_para
        new_para = _insert_para_after(insert_after, options['project_type'].upper(), size=11, small_caps=True)
        insert_after = new_para
        additions += 2

    # Degree info
    if options.get('course'):
        new_para = _insert_para_after(insert_after, '')
        insert_after = new_para
        new_para = _insert_para_after(insert_after, 'submitted in partial fulfillment of the requirements for the degree of', size=10)
        insert_after = new_para
        new_para = _insert_para_after(insert_after, options['course'], size=12, bold=True)
        insert_after = new_para
        additions += 3

    # Student ID
    if options.get('student_id'):
        new_para = _insert_para_after(insert_after, '')
        insert_after = new_para
        new_para = _insert_para_after(insert_after, f"Matriculation No.: {options['student_id']}", size=10)
        insert_after = new_para
        additions += 2

    # Supervisors
    if options.get('instructor'):
        new_para = _insert_para_after(insert_after, '')
        insert_after = new_para
        new_para = _insert_para_after(insert_after, f"First Supervisor: {options['instructor']}", size=10)
        insert_after = new_para
        additions += 2

    if options.get('second_examiner'):
        new_para = _insert_para_after(insert_after, f"Second Examiner: {options['second_examiner']}", size=10)
        insert_after = new_para
        additions += 1

    # System credit
    if options.get('system_credit'):
        new_para = _insert_para_after(insert_after, '')
        insert_after = new_para
        new_para = _insert_para_after(insert_after, options['system_credit'], size=9, italic=True)
        insert_after = new_para
        additions += 2

    # Location
    if options.get('location'):
        new_para = _insert_para_after(insert_after, '')
        insert_after = new_para
        new_para = _insert_para_after(insert_after, options['location'], size=10)
        insert_after = new_para
        additions += 2

    if verbose and additions > 0:
        print(f"   ✓ Added metadata ({additions} lines)")


def _insert_para_after(after_para, text: str, size: int = 11, bold: bool = False,
                       italic: bool = False, small_caps: bool = False):
    """Insert a new centered paragraph after the given paragraph."""
    # Get the parent element and find position
    parent = after_para._element.getparent()
    index = list(parent).index(after_para._element)

    # Create new paragraph
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    new_p = OxmlElement('w:p')

    # Add paragraph properties for centering
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    new_p.append(pPr)

    # Add run with text
    if text:
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Font size
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size * 2))  # Half-points
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(size * 2))
        rPr.append(szCs)

        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        if italic:
            i = OxmlElement('w:i')
            rPr.append(i)
        if small_caps:
            sc = OxmlElement('w:smallCaps')
            rPr.append(sc)

        run.append(rPr)

        t = OxmlElement('w:t')
        t.text = text
        run.append(t)
        new_p.append(run)

    # Insert after the target paragraph
    parent.insert(index + 1, new_p)

    # Return a paragraph object for the new element
    # Find it in the document
    doc = after_para._element.getparent().getparent()
    for para in after_para._element.getparent().iterchildren(qn('w:p')):
        if para is new_p:
            from docx.text.paragraph import Paragraph
            return Paragraph(new_p, after_para._parent)

    return after_para  # Fallback


def _insert_page_break_after(doc: Document, para_index: int) -> None:
    """Insert a page break after the specified paragraph."""
    if para_index is None or para_index >= len(doc.paragraphs):
        return

    target_para = doc.paragraphs[para_index]
    run = target_para.add_run()
    run.add_break(WD_BREAK.PAGE)


def _fix_table_widths(doc: Document, verbose: bool = False):
    """
    Fix table widths to fit within page margins.

    Pandoc-generated tables often overflow. This function:
    - Sets tables to auto-fit contents
    - Reduces font size for wide tables
    - Prevents mid-word breaks by setting noWrap on first column
    - Uses autofit layout to distribute column widths naturally
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tables_fixed = 0

    for table in doc.tables:
        num_cols = len(table.columns)

        # Get table properties
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Set table width to 100% (5000 = 100% in fifths of a percent)
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')

        # Set layout to autofit (not fixed) - allows columns to resize
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = OxmlElement('w:tblLayout')
            tblPr.append(tblLayout)
        tblLayout.set(qn('w:type'), 'autofit')

        # Enable autofit
        table.autofit = True

        # Set font size based on column count
        if num_cols >= 5:
            font_size = Pt(8)
        elif num_cols >= 4:
            font_size = Pt(9)
        else:
            font_size = Pt(10)

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.tcPr
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)

                # Remove fixed cell width constraints
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)

                # For the first column (often labels like "Domain"), prevent word wrap
                # This stops mid-word breaks like "Dermat-ology"
                if col_idx == 0:
                    noWrap = tcPr.find(qn('w:noWrap'))
                    if noWrap is None:
                        noWrap = OxmlElement('w:noWrap')
                        tcPr.append(noWrap)

                # Set font size for all cells
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = font_size
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after = Pt(2)

        tables_fixed += 1

    if verbose and tables_fixed > 0:
        print(f"   ✓ Fixed {tables_fixed} tables (auto-fit, {font_size.pt}pt font)")


# ============================================================================
# STANDALONE TESTING
# ============================================================================

def main():
    """Test DOCX post-processor."""
    import sys

    if len(sys.argv) > 1:
        docx_path = Path(sys.argv[1])
    else:
        docx_path = Path('examples/opensource_draft.docx')

    if not docx_path.exists():
        print(f"❌ File not found: {docx_path}")
        sys.exit(1)

    # Backup
    backup_path = docx_path.with_suffix('.docx.backup')
    import shutil
    shutil.copy2(docx_path, backup_path)
    print(f"📋 Created backup: {backup_path}")

    # Test options
    test_options = {
        'institution': 'Technical University of Berlin',
        'faculty': 'Faculty of Electrical Engineering and Computer Science',
        'department': 'Department of Security in Telecommunications',
        'course': 'Master of Science in Computer Science',
        'instructor': 'Prof. Dr. Maria Schmidt',
        'second_examiner': 'Prof. Dr. Hans Weber',
        'student_id': '123456',
        'project_type': 'Master Draft',
        'system_credit': 'Generated with OpenDraft AI',
        'location': 'Berlin',
    }

    success = insert_academic_structure(docx_path, verbose=True, options=test_options)

    if success:
        print("✅ Test completed!")
    else:
        print("❌ Test failed")
        shutil.copy2(backup_path, docx_path)


if __name__ == '__main__':
    main()
