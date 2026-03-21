#!/usr/bin/env python3
"""
ABOUTME: Production-grade academic document export utility with strategy pattern
ABOUTME: Supports multiple PDF engines (LibreOffice, Pandoc, WeasyPrint) with auto-fallback
"""

import sys
import argparse
from pathlib import Path
from typing import Optional, Literal

# Use centralized logging system
from utils.logging_config import get_logger
logger = get_logger(__name__)

# Add project root to path for subprocess/different CWD execution
project_root = Path(__file__).resolve().parent.parent
if str(project_root) not in sys.path:
    sys.path.insert(0, str(project_root))

from utils.pdf_engines import (
    PDFGenerationOptions,
    PDFEngineFactory,
    get_available_engines,
    get_recommended_engine
)


def extract_metadata_from_yaml(md_file: Path) -> dict:
    """
    Extract metadata from YAML frontmatter in markdown file.

    Normalizes German/Spanish/French field names to English.

    Args:
        md_file: Path to markdown file with YAML frontmatter

    Returns:
        dict: Normalized metadata (title, author, date, institution, department, degree)
    """
    import yaml

    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # Extract YAML frontmatter
        if not content.strip().startswith('---'):
            return {}

        parts = content.split('---', 2)
        if len(parts) < 3:
            return {}

        yaml_content = parts[1]
        metadata = yaml.safe_load(yaml_content) or {}

        # Normalize localized field names to English
        field_map = {
            # German
            'titel': 'title',
            'untertitel': 'subtitle',
            'autor': 'author',
            'datum': 'date',
            # Spanish
            'título': 'title',
            'subtítulo': 'subtitle',
            'autor': 'author',
            'fecha': 'date',
            # French
            'titre': 'title',
            'sous-titre': 'subtitle',
            'auteur': 'author',
            'date': 'date',
            # English (pass-through)
            'title': 'title',
            'subtitle': 'subtitle',
            'author': 'author',
            'date': 'date',
            'institution': 'institution',
            'department': 'department',
            'degree': 'degree',
            'advisor': 'advisor',
            'student_id': 'student_id',
            'project_type': 'project_type',
            'system_credit': 'system_credit',
            'generated_by': 'system_credit',  # Map generated_by to system_credit
        }

        normalized = {}
        for key, value in metadata.items():
            normalized_key = field_map.get(key.lower(), key.lower())
            normalized[normalized_key] = value

        return normalized

    except Exception as e:
        logger.warning(f"Could not extract YAML metadata: {e}")
        return {}


def export_pdf(
    md_file: Path,
    output_pdf: Path,
    engine: Literal['auto', 'libreoffice', 'pandoc', 'weasyprint'] = 'auto',
    options: Optional[PDFGenerationOptions] = None
) -> bool:
    """
    Export markdown to PDF using specified engine.

    Uses strategy pattern with automatic fallback to other engines on failure.
    This follows SOLID principles and provides production-grade error handling.

    Args:
        md_file: Path to input markdown file
        output_pdf: Path for output PDF file
        engine: PDF engine to use ('auto' for automatic selection)
        options: PDF generation options (uses academic defaults if None)

    Returns:
        bool: True if PDF generation succeeded, False otherwise

    Examples:
        >>> export_pdf(Path('draft.md'), Path('draft.pdf'))
        >>> export_pdf(Path('draft.md'), Path('draft.pdf'), engine='libreoffice')
        >>> options = PDFGenerationOptions(line_spacing=1.5)
        >>> export_pdf(Path('draft.md'), Path('draft.pdf'), options=options)
    """
    # Extract metadata from YAML frontmatter
    metadata = extract_metadata_from_yaml(md_file)

    # Create options with metadata if not provided
    if options is None:
        options = PDFGenerationOptions(
            title=metadata.get('title'),
            subtitle=metadata.get('subtitle'),
            author=metadata.get('author'),
            date=metadata.get('date'),
            institution=metadata.get('institution'),
            department=metadata.get('department'),
            course=metadata.get('degree'),  # Map degree to course field
            instructor=metadata.get('advisor'),  # Map advisor to instructor field
            student_id=metadata.get('student_id'),
            project_type=metadata.get('project_type'),
            system_credit=metadata.get('system_credit')
        )

    logger.info("="*70)
    logger.info(f"Generating PDF: {output_pdf.name}")
    logger.info(f"Input: {md_file}")
    logger.info(f"Engine: {engine}")
    logger.info("="*70)

    # Use factory to generate with automatic fallback
    result = PDFEngineFactory.generate_with_fallback(
        md_file=md_file,
        output_pdf=output_pdf,
        options=options,
        preferred_engine=engine if engine != 'auto' else None
    )

    # Display result
    if result.success:
        logger.info(f"PDF created successfully: {result.output_path}")
        logger.info(f"Engine used: {result.engine_name}")

        if result.warnings:
            logger.warning("Warnings during PDF generation:")
            for warning in result.warnings:
                logger.warning(f"  - {warning}")

        return True
    else:
        logger.error("PDF generation failed")
        logger.error(f"Error: {result.error_message}")
        return False


def _normalize_yaml_for_pandoc(md_content: str) -> str:
    """
    Normalize YAML frontmatter field names to English for Pandoc compatibility.

    Pandoc only recognizes English field names (title, subtitle, author, date)
    for creating styled paragraphs in DOCX/PDF output. This function translates
    localized field names back to English while preserving the field values.

    Supported translations:
    - German: titel→title, untertitel→subtitle, autor→author, datum→date
    - Spanish: título→title, subtítulo→subtitle, autor→author, fecha→date
    - French: titre→title, sous-titre→subtitle, auteur→author, date→date

    Args:
        md_content: Markdown content with YAML frontmatter

    Returns:
        Markdown content with normalized YAML field names
    """
    import re

    # Translation map: localized → English
    field_translations = {
        # German
        'titel:': 'title:',
        'untertitel:': 'subtitle:',
        'autor:': 'author:',
        'datum:': 'date:',
        'wortzahl:': 'word_count:',
        'seitenzahl:': 'page_count:',
        'sprache:': 'language:',
        'thema:': 'topic:',
        'schlagwörter:': 'keywords:',
        'qualitäts_bewertung:': 'quality_score:',
        'system_ersteller:': 'system_creator:',
        'zitate_verifiziert:': 'citations_verified:',
        'visuelle_elemente:': 'visual_elements:',
        'generierungs_methode:': 'generation_method:',
        'beschreibung_showcase:': 'showcase_description:',
        'system_fähigkeiten:': 'system_capabilities:',
        'aufruf_zur_aktion:': 'call_to_action:',
        'lizenz:': 'license:',

        # Spanish
        'título:': 'title:',
        'subtítulo:': 'subtitle:',
        'fecha:': 'date:',
        'recuento_de_palabras:': 'word_count:',
        'idioma:': 'language:',

        # French
        'titre:': 'title:',
        'sous-titre:': 'subtitle:',
        'auteur:': 'author:',
        'nombre_de_mots:': 'word_count:',
        'langue:': 'language:',
    }

    # Only process if YAML frontmatter exists
    if not md_content.strip().startswith('---'):
        return md_content

    # Extract YAML frontmatter
    parts = md_content.split('---', 2)
    if len(parts) < 3:
        return md_content

    yaml_content = parts[1]
    rest_content = parts[2]

    # Translate field names (case-insensitive)
    for localized, english in field_translations.items():
        yaml_content = re.sub(
            f'^{re.escape(localized)}',
            english,
            yaml_content,
            flags=re.MULTILINE | re.IGNORECASE
        )

    # Reconstruct markdown
    return f'---{yaml_content}---{rest_content}'


def export_docx_basic(md_file: Path, output_docx: Path) -> bool:
    """
    Export markdown to DOCX format using basic python-docx parsing.

    Note: This is a legacy/fallback function with limited formatting support.
    Does NOT support tables, complex formatting, or citations.

    For production use, prefer export_docx() which uses Pandoc with full
    markdown support and proper table formatting.

    Args:
        md_file: Path to input markdown file
        output_docx: Path for output DOCX file

    Returns:
        bool: True if DOCX generation succeeded, False otherwise
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx not installed")
        logger.error("Run: pip install python-docx")
        return False

    try:
        # Read markdown
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Normalize YAML for Pandoc (translate localized field names to English)
        md_content = _normalize_yaml_for_pandoc(md_content)
        lines = md_content.splitlines(keepends=True)

        # Create document
        doc = Document()

        # Set margins (1 inch)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Add footer instruction for page numbers
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run("[Add page numbers via Insert > Page Number in Word]")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.italic = True

        # Process markdown content
        for line in lines:
            line = line.rstrip()

            if not line:
                continue

            # Title (# heading)
            if line.startswith('# '):
                para = doc.add_heading(line[2:], level=1)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            # Section heading (## heading)
            if line.startswith('## '):
                doc.add_heading(line[3:], level=2)
                continue

            # Subsection heading (### heading)
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
                continue

            # Horizontal rule
            if line.startswith('---'):
                para = doc.add_paragraph()
                run = para.add_run('_' * 60)
                run.font.size = Pt(12)
                continue

            # Regular paragraph
            para = doc.add_paragraph(line)
            para_format = para.paragraph_format
            para_format.line_spacing = 2.0  # Double spacing
            para_format.space_after = Pt(0)

            # Set font
            if para.runs:
                run = para.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        # Save
        doc.save(output_docx)

        logger.info(f"DOCX created successfully: {output_docx}")
        return True

    except Exception as e:
        logger.error(f"DOCX generation failed: {str(e)}")
        return False


def export_docx(
    md_file: Path,
    output_docx: Path,
    options: Optional[PDFGenerationOptions] = None
) -> bool:
    """
    Export markdown to DOCX with full formatting support (tables, citations, styles).

    Uses Pandoc with custom reference document for professional formatting that matches
    PDF output quality. Automatically falls back to basic export if Pandoc unavailable.

    Features:
    - ✅ Tables rendered as Word tables (not pipe text)
    - ✅ Proper heading styles (APA 7th edition compatible)
    - ✅ Citations and footnotes
    - ✅ Table of contents
    - ✅ Custom fonts and spacing from reference document

    Args:
        md_file: Path to input markdown file
        output_docx: Path for output DOCX file
        options: PDF generation options (uses academic defaults if None)

    Returns:
        bool: True if DOCX generation succeeded, False otherwise

    Examples:
        >>> export_docx(Path('draft.md'), Path('draft.docx'))
        >>> options = PDFGenerationOptions(enable_toc=True, toc_depth=3)
        >>> export_docx(Path('draft.md'), Path('draft.docx'), options)
    """
    import subprocess
    import shutil

    # Extract metadata from YAML frontmatter
    metadata = extract_metadata_from_yaml(md_file)

    # Create options with metadata if not provided
    if options is None:
        options = PDFGenerationOptions(
            title=metadata.get('title'),
            subtitle=metadata.get('subtitle'),
            author=metadata.get('author'),
            date=metadata.get('date'),
            institution=metadata.get('institution'),
            department=metadata.get('department'),
            faculty=metadata.get('faculty'),
            course=metadata.get('degree'),  # Map degree to course field
            instructor=metadata.get('advisor'),  # Map advisor to instructor field
            second_examiner=metadata.get('second_examiner'),
            student_id=metadata.get('student_id'),
            project_type=metadata.get('project_type'),
            system_credit=metadata.get('system_credit'),
            location=metadata.get('location')
        )

    # Try Pandoc method first (best quality)
    if not shutil.which('pandoc'):
        logger.warning("Pandoc not found - falling back to basic DOCX export")
        logger.warning("Tables and advanced formatting will be limited")
        logger.info("Install Pandoc for better results: sudo apt install pandoc")
        return export_docx_basic(md_file, output_docx)

    try:
        # Read and normalize YAML field names for Pandoc compatibility
        # (Pandoc only recognizes English field names like 'title', 'author', 'date')
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        md_content = _normalize_yaml_for_pandoc(md_content)

        # Write normalized content to temporary file for Pandoc
        import tempfile
        temp_md = None
        try:
            temp_fd, temp_path = tempfile.mkstemp(suffix='.md', text=True)
            temp_md = Path(temp_path)
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(md_content)
        finally:
            if temp_fd:
                import os
                os.close(temp_fd)

        # Locate reference document
        reference_doc = Path(__file__).parent.parent / "examples" / "custom-reference.docx"

        if not reference_doc.exists():
            logger.warning(f"Reference document not found: {reference_doc}")
            logger.warning("Continuing without reference document (may lose some formatting)")
            reference_doc = None

        # Build pandoc command (use normalized temp file)
        cmd = [
            'pandoc',
            str(temp_md),  # Use normalized markdown instead of original
            '-o', str(output_docx),
            '--from', 'markdown',
            '--to', 'docx',
        ]

        # Add reference document for styling
        if reference_doc:
            cmd.extend(['--reference-doc', str(reference_doc)])

        # Add table of contents (Pandoc generates a proper Word TOC field)
        if options.enable_toc:
            cmd.append('--toc')
            cmd.extend(['--toc-depth', str(options.toc_depth)])

        # Add metadata if provided
        if options.title:
            cmd.extend(['--metadata', f'title={options.title}'])
        if options.author:
            cmd.extend(['--metadata', f'author={options.author}'])

        logger.info("="*70)
        logger.info(f"Generating DOCX with Pandoc: {output_docx.name}")
        logger.info(f"Input: {md_file}")
        if reference_doc:
            logger.info(f"Reference doc: {reference_doc.name}")
        logger.info("="*70)

        # Run pandoc
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60
        )

        if result.returncode != 0:
            logger.error(f"Pandoc failed with return code {result.returncode}")
            if result.stderr:
                logger.error(f"Error: {result.stderr}")
            return False

        logger.info(f"DOCX created successfully: {output_docx}")
        logger.info("Tables, formatting, and styling preserved from markdown")

        # Post-process DOCX to add academic structure (title page + TOC + page breaks)
        # Fixes Pandoc's inline title block by inserting professional page breaks
        from utils.docx_post_processor import insert_academic_structure

        # Build options dict from PDFGenerationOptions for cover page enhancement
        post_options = {}
        if options:
            if options.institution:
                post_options['institution'] = options.institution
            if hasattr(options, 'faculty') and options.faculty:
                post_options['faculty'] = options.faculty
            if options.department:
                post_options['department'] = options.department
            if options.course:
                post_options['course'] = options.course
            if options.instructor:
                post_options['instructor'] = options.instructor
            if hasattr(options, 'second_examiner') and options.second_examiner:
                post_options['second_examiner'] = options.second_examiner
            if options.student_id:
                post_options['student_id'] = options.student_id
            if hasattr(options, 'project_type') and options.project_type:
                post_options['project_type'] = options.project_type
            if hasattr(options, 'system_credit') and options.system_credit:
                post_options['system_credit'] = options.system_credit
            if hasattr(options, 'location') and options.location:
                post_options['location'] = options.location

        if not insert_academic_structure(output_docx, verbose=True, options=post_options if post_options else None):
            logger.warning("Post-processing failed - DOCX created but may lack page structure")
            logger.warning("DOCX will have inline title block instead of standalone pages")
            return True  # Still return True since basic DOCX was created

        return True

    except subprocess.TimeoutExpired:
        logger.error("DOCX generation timed out (>60s)")
        return False
    except Exception as e:
        logger.error(f"DOCX generation failed: {str(e)}")
        return False
    finally:
        # Clean up temporary markdown file
        if temp_md and temp_md.exists():
            temp_md.unlink()


def show_available_engines() -> None:
    """Display available PDF engines and their status."""
    logger.info("="*70)
    logger.info("Available PDF Engines")
    logger.info("="*70)

    engines = get_available_engines()
    if not engines:
        logger.warning("No PDF engines available")
        logger.info("Install at least one of:")
        logger.info("  - LibreOffice: sudo apt install libreoffice-writer libreoffice-core-nogui")
        logger.info("  - Pandoc: sudo apt install pandoc texlive-latex-base texlive-latex-recommended")
        logger.info("  - WeasyPrint: pip install weasyprint")
        return

    recommended = get_recommended_engine()

    for engine in engines:
        marker = "[Recommended]" if engine == recommended else ""
        logger.info(f"  {engine} {marker}")

    if recommended:
        logger.info(f"Recommended: {recommended}")


def main():
    """CLI entry point."""
    # Logging is automatically configured by utils.logging_config on import

    parser = argparse.ArgumentParser(
        description='Export markdown to professional academic PDF/DOCX',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  # Auto-select best engine
  python export_professional.py draft.md --pdf draft.pdf

  # Use specific engine
  python export_professional.py draft.md --pdf draft.pdf --engine libreoffice

  # Generate both PDF and DOCX
  python export_professional.py draft.md --pdf draft.pdf --docx draft.docx

  # List available engines
  python export_professional.py --list-engines

Engine Options:
  auto         - Automatically select best available engine (recommended)
  libreoffice  - Use LibreOffice (best quality, requires installation)
  pandoc       - Use Pandoc/LaTeX (highest quality, academic standard)
  weasyprint   - Use WeasyPrint (fallback, known font rendering issues)
        '''
    )
    parser.add_argument('input', nargs='?', help='Input markdown file')
    parser.add_argument('--pdf', help='Output PDF file')
    parser.add_argument('--docx', help='Output DOCX file')
    parser.add_argument(
        '--engine',
        choices=['auto', 'libreoffice', 'pandoc', 'weasyprint'],
        default='auto',
        help='PDF engine to use (default: auto)'
    )
    parser.add_argument(
        '--list-engines',
        action='store_true',
        help='List available PDF engines and exit'
    )

    args = parser.parse_args()

    # List engines mode
    if args.list_engines:
        show_available_engines()
        sys.exit(0)

    # Validate arguments
    if not args.input:
        parser.error("Input file required (or use --list-engines)")

    input_file = Path(args.input)

    if not input_file.exists():
        logger.error(f"Input file not found: {input_file}")
        sys.exit(1)

    if not args.pdf and not args.docx:
        parser.error("Specify --pdf and/or --docx output file")

    # Generate outputs
    success = True

    if args.pdf:
        pdf_success = export_pdf(input_file, Path(args.pdf), engine=args.engine)
        success = success and pdf_success

    if args.docx:
        docx_success = export_docx(input_file, Path(args.docx))
        success = success and docx_success

    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
